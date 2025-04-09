from docx import Document
from datetime import datetime


class Reporter:
    def __init__(self, articles):
        #self.title = title
        self.date = datetime.now().strftime("%Y-%m-%d")
        self.master_title = f"{self.date} - News Report"
        self.articles = articles
    
    def dedupe_extracted_information(self, important_info):
        deduped_info = {}
        for key, value in important_info.items():
            deduped_info[key] = list(dict.fromkeys(value))
        return deduped_info
        
    def generate_report(self):
        document = Document()
        
        document.add_heading(self.master_title, level=0)
        
        for title, content in self.articles.items():
            document.add_heading(title, level=1)
            
            
            
            summary = content.get("Summary", "No summary available.")
            #original_content = content.get("Original Content", "No original content available.")
            document.add_paragraph(f"Summary:\n {summary}")
            #document.add_paragraph(f"original contnet: {original_content}")
            
            important = content.get("Important Info",{})
            if important:
                document.add_paragraph("Extracted Important Information:")
                for key, value in self.dedupe_extracted_information(important).items():
                    for val in value:
                        document.add_paragraph(f"{key}: {val}", style="List Bullet")
                          
            if "link" in content:
                document.add_paragraph(f"Link: {content["link"]}")

        filename = f"{self.master_title}.docx"
        document.save(filename)
        print(f"Report saved as \033[1,32m{filename}\033[0m")
        